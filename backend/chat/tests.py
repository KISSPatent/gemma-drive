"""
Tests for chat.extraction.

Drive service calls are mocked so these tests don't touch the network
or require OAuth. python-docx is used directly to build an in-memory
DOCX so the DOCX extraction path is exercised end-to-end.
"""
import io
from unittest.mock import patch, MagicMock

from django.test import TestCase
from docx import Document

from chat.extraction import (
    extract,
    SUPPORTED_MIMES,
    MAX_CHARS,
    PLAIN_TEXT,
    DOCX,
    PDF,
)


class ExtractionTests(TestCase):
    def test_plain_text_round_trip(self):
        """Plain text files come back unmodified."""
        service = MagicMock()
        with patch("chat.extraction._download_bytes", return_value=b"hello world"):
            text, truncated = extract(service, "fid-1", PLAIN_TEXT, "notes.txt")
        self.assertEqual(text, "hello world")
        self.assertFalse(truncated)

    def test_docx_extraction(self):
        """Real DOCX bytes are parsed and paragraphs are concatenated."""
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        buf = io.BytesIO()
        doc.save(buf)

        service = MagicMock()
        with patch("chat.extraction._download_bytes", return_value=buf.getvalue()):
            text, truncated = extract(service, "fid-2", DOCX, "doc.docx")

        self.assertIn("First paragraph", text)
        self.assertIn("Second paragraph", text)
        self.assertFalse(truncated)

    def test_unsupported_mime_raises(self):
        """Unknown MIME types raise ValueError, not a silent empty extraction."""
        service = MagicMock()
        with self.assertRaises(ValueError):
            extract(service, "fid-3", "application/x-not-a-real-type", "weird.bin")

    def test_truncation_at_max_chars(self):
        """Text over MAX_CHARS is truncated and the flag is set."""
        oversize = b"x" * (MAX_CHARS + 100)
        service = MagicMock()
        with patch("chat.extraction._download_bytes", return_value=oversize):
            text, truncated = extract(service, "fid-4", PLAIN_TEXT, "big.txt")

        self.assertTrue(truncated)
        self.assertTrue(text.startswith("x" * MAX_CHARS))
        self.assertIn("truncated", text)

    def test_supported_mimes_includes_core_formats(self):
        """SUPPORTED_MIMES is the source of truth for the UI's 'supported' flag,
        so guard against accidental removal."""
        for required in (PDF, DOCX, PLAIN_TEXT):
            self.assertIn(required, SUPPORTED_MIMES)
